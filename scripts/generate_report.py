# -*- coding: utf-8 -*-
"""
Báo cáo đồ án môn học FoodNow — Nhóm 9 (v2)
Viết lại theo mẫu chuẩn BaoCao_DoAn_HDNK.docx

Chạy: python scripts/generate_report.py
Output: BaoCaoDoAn_FoodNow_Nhom9.docx (thư mục gốc dự án)
"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os


# ═══════════════════════════════════════════════════════════════════
# HELPERS
# ═══════════════════════════════════════════════════════════════════

def setup_styles(doc):
    """Cấu hình font/spacing các built-in style theo mẫu chuẩn."""
    # Normal: Times New Roman 13pt
    n = doc.styles['Normal']
    n.font.name = 'Times New Roman'
    n.font.size = Pt(13)

    # Heading 1 — tiêu đề chương (căn giữa, 14pt, đậm)
    h1s = doc.styles['Heading 1']
    h1s.font.name = 'Times New Roman'
    h1s.font.size = Pt(14)
    h1s.font.bold = True
    h1s.font.italic = False
    h1s.font.color.rgb = RGBColor(0, 0, 0)
    h1s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    h1s.paragraph_format.space_before = Pt(18)
    h1s.paragraph_format.space_after = Pt(4)

    # Heading 2 — tiêu đề mục (13pt, đậm)
    h2s = doc.styles['Heading 2']
    h2s.font.name = 'Times New Roman'
    h2s.font.size = Pt(13)
    h2s.font.bold = True
    h2s.font.italic = False
    h2s.font.color.rgb = RGBColor(0, 0, 0)
    h2s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h2s.paragraph_format.space_before = Pt(8)
    h2s.paragraph_format.space_after = Pt(4)

    # Heading 3 — tiêu đề tiểu mục (13pt, đậm nghiêng)
    h3s = doc.styles['Heading 3']
    h3s.font.name = 'Times New Roman'
    h3s.font.size = Pt(13)
    h3s.font.bold = True
    h3s.font.italic = True
    h3s.font.color.rgb = RGBColor(0, 0, 0)
    h3s.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    h3s.paragraph_format.space_before = Pt(8)
    h3s.paragraph_format.space_after = Pt(4)

    # List Paragraph
    lp = doc.styles['List Paragraph']
    lp.font.name = 'Times New Roman'
    lp.font.size = Pt(13)
    lp.paragraph_format.left_indent = Cm(1.0)
    lp.paragraph_format.space_before = Pt(1)
    lp.paragraph_format.space_after = Pt(2)


def h1(doc, text):
    return doc.add_paragraph(text, style='Heading 1')


def h2(doc, text):
    return doc.add_paragraph(text, style='Heading 2')


def h3(doc, text):
    return doc.add_paragraph(text, style='Heading 3')


def body(doc, text='', indent=True, align=WD_ALIGN_PARAGRAPH.LEFT,
         space_before=0, space_after=6, bold=False, italic=False):
    p = doc.add_paragraph(style='Normal')
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
    return p


def bullet(doc, text, level=0):
    p = doc.add_paragraph(style='List Paragraph')
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.5)
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    p.add_run('– ' + text)
    return p


def spacer(doc, space_after=6):
    p = doc.add_paragraph(style='Normal')
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    return p


def page_break(doc):
    doc.add_page_break()


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def fill_cell(cell, text, bold=False, italic=False,
              align=WD_ALIGN_PARAGRAPH.LEFT, size=12):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size)
        run.font.name = 'Times New Roman'


def make_table(doc, headers, data_rows, col_widths=None, header_bg='365F91'):
    """Tạo bảng chuẩn với header màu xanh và dữ liệu bên dưới."""
    n_cols = len(headers)
    t = doc.add_table(rows=1 + len(data_rows), cols=n_cols)
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        set_cell_bg(c, header_bg)
        fill_cell(c, h, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    for ri, row_data in enumerate(data_rows):
        for ci, val in enumerate(row_data):
            c = t.rows[ri + 1].cells[ci]
            if isinstance(val, tuple):
                txt, al = val
            else:
                txt, al = val, WD_ALIGN_PARAGRAPH.LEFT
            fill_cell(c, txt, align=al, size=12)
    if col_widths:
        for row in t.rows:
            for i, w in enumerate(col_widths):
                if i < len(row.cells):
                    row.cells[i].width = Cm(w)
    spacer(doc, 8)
    return t


# ═══════════════════════════════════════════════════════════════════
# TRANG BÌA
# ═══════════════════════════════════════════════════════════════════

def build_cover(doc):
    def cp(text='', bold=False, size=13, space_after=2, italic=False):
        p = doc.add_paragraph(style='Normal')
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            r = p.add_run(text)
            r.bold = bold
            r.italic = italic
            if size != 13:
                r.font.size = Pt(size)
        return p

    cp('TRƯỜNG ĐẠI HỌC ĐÀ LẠT', bold=True)
    cp('KHOA CÔNG NGHỆ THÔNG TIN', bold=True, space_after=16)
    cp('─────────────────────────────────', space_after=24)
    cp('BÁO CÁO ĐỒ ÁN MÔN HỌC', bold=True, size=16, space_after=4)
    cp('Môn học: Phát Triển Ứng Dụng Android', bold=True, size=14, space_after=4)
    cp('Học kỳ 2, năm học 2025 – 2026', italic=True, size=13, space_after=32)
    cp('ĐỀ TÀI:', bold=True, size=14, space_after=6)
    cp('FOODNOW', bold=True, size=20, space_after=4)
    cp('ỨNG DỤNG ĐẶT ĐỒ ĂN TRỰC TUYẾN TRÊN NỀN TẢNG ANDROID',
       bold=True, size=14, space_after=40)
    cp('GIẢNG VIÊN HƯỚNG DẪN:', bold=True, space_after=4)
    cp('TS. Nguyễn Thị Lương', space_after=2)
    cp('KS. La Quốc Thắng', space_after=30)
    cp('NHÓM THỰC HIỆN — NHÓM 9', bold=True, space_after=6)
    for mssv, name in [('2312660', 'Đinh Thị Mai Lành'),
                       ('2312567', 'Võ Thị Minh Ân'),
                       ('2312708', 'Trương Võ Trọng Nhân')]:
        cp(f'{mssv}  –  {name}', space_after=2)
    cp(space_after=36)
    cp('Đà Lạt, tháng 4 năm 2026', italic=True)
    page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# NHẬN XÉT GIẢNG VIÊN
# ═══════════════════════════════════════════════════════════════════

def build_nhan_xet(doc):
    p = doc.add_paragraph('NHẬN XÉT CỦA GIẢNG VIÊN HƯỚNG DẪN', style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)

    body(doc, '─' * 55, indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)

    for _ in range(20):
        body(doc, '', indent=False, space_after=12)

    sig = doc.add_paragraph(style='Normal')
    sig.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig.paragraph_format.space_after = Pt(6)
    sig.add_run('Đà Lạt, ngày     tháng     năm 2026')

    sig2 = doc.add_paragraph(style='Normal')
    sig2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig2.paragraph_format.space_after = Pt(4)
    r = sig2.add_run('Giảng viên hướng dẫn')
    r.bold = True

    for _ in range(5):
        body(doc, '', indent=False, space_after=12)

    sig3 = doc.add_paragraph(style='Normal')
    sig3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    sig3.paragraph_format.space_after = Pt(4)
    sig3.add_run('(Ký và ghi rõ họ tên)').italic = True

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# LỜI CẢM ƠN
# ═══════════════════════════════════════════════════════════════════

def build_loi_cam_on(doc):
    p = doc.add_paragraph('LỜI CẢM ƠN', style='Normal')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(16)
    p.runs[0].bold = True
    p.runs[0].font.size = Pt(14)

    body(doc,
         'Đầu tiên, nhóm chúng em xin gửi lời cảm ơn chân thành và sâu sắc nhất đến '
         'TS. Nguyễn Thị Lương và KS. La Quốc Thắng — những người đã tận tình hướng '
         'dẫn, giải đáp thắc mắc và động viên nhóm trong suốt quá trình thực hiện đồ '
         'án. Sự định hướng của thầy cô đã giúp nhóm nắm bắt được kiến thức và phương '
         'pháp tiếp cận đúng đắn khi xây dựng ứng dụng Android theo kiến trúc MVVM.',
         space_after=8)

    body(doc,
         'Chúng em cũng xin gửi lời cảm ơn đến Khoa Công nghệ Thông tin — Trường Đại '
         'học Đà Lạt đã tạo điều kiện thuận lợi để chúng em được học tập, tiếp cận '
         'với các công nghệ hiện đại và có cơ hội thực hành qua những đồ án học phần '
         'thực tiễn như môn Phát Triển Ứng Dụng Android này.',
         space_after=8)

    body(doc,
         'Tuy nhiên, do kiến thức và kinh nghiệm còn hạn chế, báo cáo và sản phẩm '
         'của nhóm khó tránh khỏi những thiếu sót. Nhóm rất mong nhận được sự chỉ '
         'dẫn và góp ý của thầy cô để nhóm có thể hoàn thiện hơn.',
         space_after=8)

    body(doc, 'Nhóm sinh viên xin chân thành cảm ơn.', space_after=24)

    p_sign = doc.add_paragraph(style='Normal')
    p_sign.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_sign.add_run('Đà Lạt, tháng 4 năm 2026').italic = True

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# ĐỀ CƯƠNG THỰC HIỆN ĐỒ ÁN
# ═══════════════════════════════════════════════════════════════════

def build_de_cuong(doc):
    def dcp(text='', bold=False, size=13, space_after=2,
            align=WD_ALIGN_PARAGRAPH.LEFT):
        p = doc.add_paragraph(style='Normal')
        p.alignment = align
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(space_after)
        if text:
            r = p.add_run(text)
            r.bold = bold
            if size != 13:
                r.font.size = Pt(size)
        return p

    dcp('Trường Đại học Đà Lạt', bold=True)
    dcp('Khoa Công nghệ Thông tin')
    dcp('─────────────────────', space_after=8)
    dcp('ĐỀ CƯƠNG THỰC HIỆN ĐỒ ÁN', bold=True, size=14,
        align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12)

    dcp('Tên đề tài: FoodNow — Ứng dụng đặt đồ ăn trực tuyến trên nền tảng Android',
        space_after=4)
    dcp('Sinh viên thực hiện:', space_after=2)
    for mssv, name in [('2312660', 'Đinh Thị Mai Lành'),
                       ('2312567', 'Võ Thị Minh Ân'),
                       ('2312708', 'Trương Võ Trọng Nhân')]:
        dcp(f'    {name}  –  MSSV: {mssv}  –  Lớp: CTK47A', space_after=2)
    dcp('Giảng viên hướng dẫn: TS. Nguyễn Thị Lương, KS. La Quốc Thắng',
        space_after=10)

    dcp('1. TỔNG QUAN ĐỀ TÀI', bold=True, space_after=4)
    dcp('FoodNow là một ứng dụng đặt đồ ăn trực tuyến chạy trên nền tảng Android, '
        'được xây dựng nhằm kết nối khách hàng với các quán ăn tại địa phương. '
        'Ứng dụng hỗ trợ hai nhóm người dùng: khách hàng (duyệt quán, đặt món, '
        'theo dõi đơn hàng) và chủ quán (quản lý thực đơn, xử lý đơn, thống kê '
        'doanh thu). Đề tài được phát triển bằng Java theo kiến trúc MVVM với '
        'Firebase làm backend.', space_after=8)

    dcp('2. MỤC TIÊU ĐỀ TÀI', bold=True, space_after=4)
    dcp('Đề tài nhằm xây dựng một ứng dụng Android hoàn chỉnh vận dụng kiến thức '
        'môn học, bao gồm: kiến trúc MVVM, LiveData, RecyclerView, Firebase '
        'Authentication + Cloud Firestore, lưu trữ và hiển thị ảnh qua Cloudinary '
        'và Glide.',
        space_after=8)

    dcp('3. NỘI DUNG ĐỀ TÀI', bold=True, space_after=4)
    for item in [
        'Phân tích và làm rõ yêu cầu đề tài',
        'Thiết kế cơ sở dữ liệu Firestore (6 collection)',
        'Thiết kế giao diện và luồng điều hướng màn hình (wireframe trên Figma)',
        'Xây dựng ứng dụng Android (modules: Auth, Home, Cart, Orders, Favorites, StoreOwner)',
        'Kiểm thử và đánh giá hệ thống',
        'Viết báo cáo tổng kết đề tài',
    ]:
        p = doc.add_paragraph(style='List Paragraph')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.add_run('– ' + item)
    spacer(doc, 8)

    dcp('4. PHẦN MỀM VÀ CÔNG CỤ SỬ DỤNG', bold=True, space_after=4)
    for item in [
        'Thiết kế giao diện: Figma.',
        'Ngôn ngữ lập trình: Java (Android SDK).',
        'Kiến trúc ứng dụng: MVVM (Model-View-ViewModel).',
        'Backend & Database: Firebase Authentication, Cloud Firestore.',
        'Lưu trữ ảnh: Cloudinary SDK 3.0.2.',
        'Load ảnh: Glide 4.16.0.',
        'IDE: Android Studio (Electric Eel trở lên), compileSdk 35.',
        'Quản lý phiên bản: Git, GitHub.',
    ]:
        p = doc.add_paragraph(style='List Paragraph')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.add_run('– ' + item)
    spacer(doc, 8)

    dcp('5. KẾ HOẠCH THỰC HIỆN', bold=True, space_after=6)
    schedule = [
        ('1', 'Khảo sát, thu thập tài liệu, xác định yêu cầu', 'Tuần 1–2', 'Hoàn thành báo cáo đề cương'),
        ('2', 'Thiết kế cơ sở dữ liệu Firestore và kiến trúc MVVM', 'Tuần 3–4', 'Schema 6 collection, package diagram'),
        ('3', 'Xây dựng module xác thực và hồ sơ (TV1)', 'Tuần 5–6', 'LoginActivity, RegisterActivity, ProfileFragment'),
        ('4', 'Xây dựng module trang chủ & quán ăn (TV2)', 'Tuần 7–8', 'HomeFragment, StoreDetailActivity, FoodAdapter'),
        ('5', 'Xây dựng module giỏ hàng, đặt hàng & yêu thích (TV3)', 'Tuần 7–8', 'CartFragment, CheckoutActivity, FavoritesFragment'),
        ('6', 'Xây dựng module chủ quán (TV1)', 'Tuần 9–10', 'StoreOwnerActivity, Dashboard, ManageFoods, Stats'),
        ('7', 'Kiểm thử, sửa lỗi, tối ưu hiệu năng', 'Tuần 11–12', 'Toàn nhóm'),
        ('8', 'Viết báo cáo và tổng kết đề tài', 'Tuần 13–14', 'Nộp báo cáo hoàn chỉnh'),
    ]
    t = doc.add_table(rows=1 + len(schedule), cols=4)
    t.style = 'Table Grid'
    for i, h in enumerate(['STT', 'Công việc', 'Thời gian', 'Ghi chú']):
        set_cell_bg(t.rows[0].cells[i], '365F91')
        fill_cell(t.rows[0].cells[i], h, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    for ri, (stt, cv, tg, ghu) in enumerate(schedule):
        row = t.rows[ri + 1]
        for ci, val in enumerate([stt, cv, tg, ghu]):
            fill_cell(row.cells[ci], val, size=12)
    for row in t.rows:
        row.cells[0].width = Cm(0.8)
        row.cells[1].width = Cm(7.5)
        row.cells[2].width = Cm(2.3)
        row.cells[3].width = Cm(4.9)
    spacer(doc, 8)

    dcp('6. DỰ KIẾN KẾT QUẢ', bold=True, space_after=4)
    for item in [
        'Ứng dụng Android FoodNow hoàn chỉnh với đầy đủ chức năng đặt hàng và quản lý quán.',
        'Giao diện trực quan theo Material Design 3, hỗ trợ Android 7.0 (API 24) trở lên.',
        'Khách hàng có thể duyệt quán, đặt món và theo dõi đơn hàng theo thời gian thực.',
        'Chủ quán có thể quản lý thực đơn, xử lý đơn và xem thống kê doanh thu.',
        'Báo cáo đồ án hoàn chỉnh, có giá trị tham khảo cho các khóa sau.',
    ]:
        p = doc.add_paragraph(style='List Paragraph')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.add_run('– ' + item)
    spacer(doc, 8)

    dcp('7. TÀI LIỆU THAM KHẢO', bold=True, space_after=4)
    for i, ref in enumerate([
        'Android Developer Documentation: https://developer.android.com/docs',
        'Firebase Documentation: https://firebase.google.com/docs',
        'Cloudinary Android SDK: https://cloudinary.com/documentation/android_integration',
        'Glide Image Loading Library: https://github.com/bumptech/glide',
        'Material Design 3: https://m3.material.io/',
        'GitHub Repository: https://github.com/minnhi09/FoodNow-k47pm',
    ], 1):
        dcp(f'[{i}] {ref}', space_after=2)
    spacer(doc, 14)

    # Ký tên cuối đề cương
    p_date = doc.add_paragraph(style='Normal')
    p_date.paragraph_format.space_after = Pt(4)
    p_date.add_run('Đà Lạt, ngày     tháng     năm 2026')

    p_sig = doc.add_paragraph(style='Normal')
    p_sig.paragraph_format.space_after = Pt(4)
    rb = p_sig.add_run('Giảng viên hướng dẫn')
    rb.bold = True
    p_sig.add_run('\t\t\t\t\tSinh viên thực hiện').bold = True

    for _ in range(4):
        body(doc, '', indent=False, space_after=10)

    p_ky = doc.add_paragraph(style='Normal')
    p_ky.paragraph_format.space_after = Pt(4)
    p_ky.add_run('\t\t\t\t\t(Ký và ghi rõ họ tên)').italic = True

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# MỤC LỤC (placeholder — người dùng tự cập nhật trong Word)
# ═══════════════════════════════════════════════════════════════════

def build_muc_luc(doc):
    h1(doc, 'MỤC LỤC')
    body(doc,
         '(Mục lục tự động — mở file trong Microsoft Word, bấm chuột phải vào '
         'đây và chọn "Update Field" để cập nhật số trang.)',
         indent=False, italic=True, space_after=12)
    page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# CHƯƠNG 1: TỔNG QUAN
# ═══════════════════════════════════════════════════════════════════

def build_chapter1(doc):
    h1(doc, 'CHƯƠNG 1: TỔNG QUAN')

    h2(doc, '1.1. Giới thiệu đề tài')
    body(doc,
         'Trong thời đại công nghệ số phát triển mạnh mẽ, việc đặt đồ ăn trực tuyến '
         'đã trở thành thói quen phổ biến của nhiều người, đặc biệt là giới trẻ và '
         'sinh viên. Các ứng dụng như GrabFood, ShopeeFood hay Baemin đã chứng minh '
         'nhu cầu thị trường là rất lớn. Tuy nhiên, các nền tảng này thường phức tạp, '
         'phụ thuộc nhiều vào hạ tầng backend lớn và không dễ tiếp cận với các quán '
         'ăn nhỏ lẻ tại địa phương.')
    body(doc,
         'Nhận thức được nhu cầu thực tiễn đó, nhóm đề xuất xây dựng ứng dụng '
         'FoodNow — một ứng dụng đặt đồ ăn Android đơn giản, gọn nhẹ, phù hợp với '
         'bối cảnh các quán ăn tại thành phố Đà Lạt. Ứng dụng được phát triển như '
         'một dự án học thuật nhằm vận dụng kiến thức môn Phát Triển Ứng Dụng Android.')
    body(doc,
         'FoodNow kết nối hai nhóm người dùng chính: Khách hàng — người có nhu cầu '
         'tìm kiếm và đặt món ăn; và Chủ quán — người quản lý thực đơn, xử lý đơn '
         'hàng và theo dõi doanh thu. Toàn bộ dữ liệu được lưu trữ và đồng bộ theo '
         'thời gian thực thông qua Cloud Firestore của Google Firebase.')

    h2(doc, '1.2. Mục tiêu đề tài')
    body(doc,
         'Mục tiêu tổng quát của đề tài là xây dựng một ứng dụng Android hoàn chỉnh, '
         'ứng dụng kiến thức lập trình di động theo kiến trúc chuẩn MVVM vào một '
         'sản phẩm có giá trị thực tiễn.')
    body(doc, 'Về mặt chức năng, hệ thống được thiết kế với các mục tiêu cụ thể:',
         space_after=4)
    for g in [
        'Xây dựng ứng dụng Android hoàn chỉnh theo kiến trúc MVVM, ngôn ngữ Java, giao diện XML.',
        'Tích hợp Firebase Authentication để xác thực người dùng (đăng ký, đăng nhập, đặt lại mật khẩu).',
        'Sử dụng Cloud Firestore làm cơ sở dữ liệu NoSQL với đồng bộ dữ liệu thời gian thực.',
        'Hỗ trợ hai vai trò người dùng: Khách hàng và Chủ quán với giao diện riêng biệt.',
        'Xây dựng đầy đủ luồng đặt hàng: duyệt quán → xem thực đơn → thêm vào giỏ → thanh toán → theo dõi đơn.',
        'Xây dựng hệ thống quản lý quán: dashboard, quản lý thực đơn CRUD, xử lý đơn hàng real-time, thống kê doanh thu.',
        'Tích hợp Cloudinary để upload và lưu trữ ảnh món ăn, ảnh quán, ảnh đại diện.',
        'Áp dụng Material Design 3 để xây dựng giao diện hiện đại, thân thiện.',
    ]:
        bullet(doc, g)
    spacer(doc, 4)

    h2(doc, '1.3. Nội dung thực hiện')
    body(doc,
         'Để đạt được các mục tiêu trên, nhóm đã triển khai thực hiện theo các bước sau:',
         space_after=4)
    for item in [
        'Tìm hiểu về công nghệ và công cụ: nghiên cứu kiến trúc MVVM, Android Jetpack '
        '(LiveData, ViewModel), Firebase, Cloudinary và các thư viện cần thiết.',
        'Phân tích yêu cầu hệ thống: xác định các tác nhân, use case và yêu cầu chức '
        'năng/phi chức năng của ứng dụng.',
        'Thiết kế hệ thống: thiết kế kiến trúc MVVM, cấu trúc package, schema cơ sở '
        'dữ liệu Firestore (6 collection), luồng điều hướng màn hình.',
        'Xây dựng ứng dụng: phát triển tuần tự các module theo phân công — module xác '
        'thực (TV1), module trang chủ & quán ăn (TV2), module giỏ hàng & yêu thích '
        '(TV3), module chủ quán (TV1).',
        'Kiểm thử và đánh giá: kiểm thử từng chức năng, sửa lỗi và tối ưu hiệu năng.',
        'Viết báo cáo và hoàn thiện tài liệu: tổng hợp toàn bộ quá trình thực hiện, '
        'kết quả đạt được và các hướng phát triển tiếp theo.',
    ]:
        p = doc.add_paragraph(style='List Paragraph')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(3)
        p.add_run('– ' + item)

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# CHƯƠNG 2: CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ
# ═══════════════════════════════════════════════════════════════════

def build_chapter2(doc):
    h1(doc, 'CHƯƠNG 2: CƠ SỞ LÝ THUYẾT VÀ CÔNG NGHỆ')

    # 2.1 MVVM
    h2(doc, '2.1 Kiến trúc MVVM')
    body(doc,
         'MVVM (Model-View-ViewModel) là một mẫu kiến trúc phần mềm được thiết kế '
         'để tách biệt giao diện người dùng (UI) khỏi logic nghiệp vụ và dữ liệu. '
         'Đây là kiến trúc được Google khuyến nghị sử dụng cho các ứng dụng Android '
         'hiện đại thông qua bộ thư viện Android Jetpack.')
    body(doc, 'MVVM gồm ba thành phần chính:', space_after=4)
    for item in [
        'Model: chứa dữ liệu và logic nghiệp vụ. Trong FoodNow, Model bao gồm các lớp '
        'Java (User, Store, Food, Order, Favorite, Category) và các Repository thực hiện '
        'tương tác với Firebase Firestore.',
        'View: là tầng giao diện người dùng, bao gồm các Activity và Fragment. View chỉ '
        'hiển thị dữ liệu và nhận sự kiện từ người dùng, không chứa logic nghiệp vụ.',
        'ViewModel: là cầu nối giữa View và Model. ViewModel chứa LiveData để thông báo '
        'thay đổi dữ liệu cho View và tồn tại qua các thay đổi cấu hình (xoay màn hình).',
    ]:
        bullet(doc, item)
    body(doc,
         'Ưu điểm chính của MVVM: tách biệt rõ ràng giữa các tầng, dễ kiểm thử, '
         'tái sử dụng code và bảo trì. Khi dữ liệu thay đổi trong ViewModel, View '
         'được tự động cập nhật thông qua cơ chế quan sát LiveData.',
         space_after=8)

    # 2.2 Nền tảng Android
    h2(doc, '2.2 Nền tảng Android')

    h3(doc, '2.2.1 Android SDK và ngôn ngữ Java')
    body(doc,
         'Android SDK (Software Development Kit) là bộ công cụ phát triển ứng dụng '
         'cho hệ điều hành Android. FoodNow sử dụng Android API Level 24 (Android 7.0) '
         'trở lên, hướng đến API Level 35 (Android 15), đảm bảo tương thích với đại '
         'đa số thiết bị Android hiện tại.')
    body(doc,
         'Ngôn ngữ lập trình Java được chọn vì tính ổn định, lượng tài liệu phong phú '
         'và phù hợp với mục tiêu học thuật của đề tài. Giao diện người dùng được xây '
         'dựng bằng XML Layout thuần (không sử dụng Jetpack Compose), kết hợp với '
         'RecyclerView, CardView và Material Design Components.')

    h3(doc, '2.2.2 Vòng đời Activity và Fragment')
    body(doc,
         'Activity và Fragment là hai thành phần giao diện cơ bản trong Android. '
         'Activity đại diện cho một màn hình hoàn chỉnh, còn Fragment là một "màn hình '
         'con" có thể được gắn vào Activity. Mỗi thành phần có vòng đời (lifecycle) '
         'riêng với các trạng thái: onCreate → onStart → onResume → onPause → onStop '
         '→ onDestroy.')
    body(doc,
         'Trong FoodNow, MainActivity là Activity chính chứa BottomNavigationView, '
         'quản lý 4 Fragment tương ứng với 4 tab: Trang chủ, Đơn hàng, Yêu thích và '
         'Tài khoản. Các màn hình phức tạp hơn như chi tiết quán ăn, thanh toán và '
         'quản lý chủ quán được triển khai dưới dạng Activity riêng.')

    # 2.3 Firebase
    h2(doc, '2.3 Firebase')
    body(doc,
         'Firebase là nền tảng Backend-as-a-Service (BaaS) của Google, cung cấp các '
         'dịch vụ backend đã được xây dựng sẵn thông qua SDK. Firebase giúp nhóm phát '
         'triển ứng dụng nhanh hơn mà không cần tự dựng server, quản lý cơ sở dữ liệu '
         'hay xây dựng hệ thống xác thực từ đầu.')

    h3(doc, '2.3.1 Firebase Authentication')
    body(doc,
         'Firebase Authentication cung cấp hệ thống xác thực người dùng đầy đủ tính '
         'năng: đăng ký bằng email/mật khẩu, đăng nhập, đăng xuất, đặt lại mật khẩu '
         'qua email. SDK tự động xử lý các vấn đề bảo mật như mã hoá mật khẩu, quản '
         'lý token JWT và phiên đăng nhập.')
    body(doc,
         'Trong FoodNow, sau khi đăng nhập thành công, Firebase trả về đối tượng '
         'FirebaseUser chứa thông tin cơ bản của người dùng. Ứng dụng sử dụng UID từ '
         'FirebaseUser làm khoá chính trong Cloud Firestore để liên kết với hồ sơ '
         'người dùng chi tiết.')

    h3(doc, '2.3.2 Cloud Firestore')
    body(doc,
         'Cloud Firestore là cơ sở dữ liệu NoSQL dạng document, lưu trữ dữ liệu theo '
         'cấu trúc Collection → Document → Fields. Một trong những tính năng nổi bật '
         'nhất là khả năng lắng nghe thay đổi dữ liệu theo thời gian thực '
         '(addSnapshotListener), rất phù hợp cho việc cập nhật trạng thái đơn hàng '
         'trong ứng dụng FoodNow.')
    body(doc,
         'Đặc điểm chính của Firestore được khai thác trong đề tài:', space_after=4)
    for f in [
        'Truy vấn linh hoạt: hỗ trợ .where(), .orderBy(), .limit() để lọc và sắp xếp.',
        'Real-time listener: addSnapshotListener() tự động cập nhật UI khi dữ liệu thay đổi.',
        'Offline support: tự động cache dữ liệu, hoạt động ngay cả khi mất kết nối mạng.',
        'Bảo mật: Firestore Security Rules kiểm soát quyền truy cập theo vai trò người dùng.',
    ]:
        bullet(doc, f)

    # 2.4 Công nghệ hỗ trợ
    h2(doc, '2.4 Công nghệ hỗ trợ')

    h3(doc, '2.4.1 Cloudinary')
    body(doc,
         'Cloudinary là dịch vụ lưu trữ và quản lý media trên đám mây. Trong FoodNow, '
         'Cloudinary được dùng để upload ảnh từ thiết bị (ảnh đại diện, ảnh quán, ảnh '
         'món ăn) và nhận về URL công khai để lưu vào Firestore. Cloudinary Android SDK '
         '(phiên bản 3.0.2) cung cấp API đơn giản để upload ảnh bất đồng bộ.')
    body(doc,
         'Lý do chọn Cloudinary thay vì Firebase Storage: Cloudinary cung cấp gói '
         'miễn phí với dung lượng đủ cho dự án học thuật, đồng thời hỗ trợ biến đổi '
         'ảnh tự động trên server (resize, crop, optimize), giúp tối ưu băng thông.')

    h3(doc, '2.4.2 Glide')
    body(doc,
         'Glide là thư viện tải và hiển thị ảnh phổ biến nhất trên Android, được '
         'phát triển bởi Bumptech. Glide xử lý tải ảnh bất đồng bộ từ URL, cache ảnh '
         'thông minh (bộ nhớ và đĩa) và hiển thị placeholder trong khi ảnh đang tải. '
         'FoodNow sử dụng Glide 4.16.0 để hiển thị ảnh quán ăn, ảnh món ăn và ảnh '
         'đại diện người dùng trong toàn bộ ứng dụng.')

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG
# ═══════════════════════════════════════════════════════════════════

def build_chapter3(doc):
    h1(doc, 'CHƯƠNG 3: PHÂN TÍCH VÀ THIẾT KẾ HỆ THỐNG')

    # 3.1
    h2(doc, '3.1 Phân tích yêu cầu chức năng')
    body(doc,
         'Hệ thống FoodNow được thiết kế để phục vụ ba nhóm tác nhân chính:')
    for actor, desc in [
        ('Khách hàng (Customer)', 'người dùng cuối có nhu cầu tìm kiếm và đặt đồ ăn.'),
        ('Chủ quán (Store Owner)', 'người quản lý quán ăn, thực đơn và đơn hàng.'),
        ('Hệ thống Firebase', 'xử lý xác thực, lưu trữ dữ liệu và gửi thông báo.'),
    ]:
        p = doc.add_paragraph(style='List Paragraph')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(actor + ': ').bold = True
        p.add_run(desc)
    spacer(doc, 4)

    body(doc, 'Bảng dưới đây tổng hợp các nhóm chức năng chính của hệ thống:', space_after=4)
    make_table(doc,
               headers=['STT', 'Nhóm chức năng', 'Mô tả', 'Tác nhân'],
               data_rows=[
                   (('1', WD_ALIGN_PARAGRAPH.CENTER), 'Xác thực người dùng',
                    'Đăng ký, đăng nhập, đăng xuất, đặt lại mật khẩu',
                    'Khách hàng, Chủ quán'),
                   (('2', WD_ALIGN_PARAGRAPH.CENTER), 'Quản lý hồ sơ',
                    'Xem và cập nhật thông tin cá nhân, ảnh đại diện',
                    'Khách hàng, Chủ quán'),
                   (('3', WD_ALIGN_PARAGRAPH.CENTER), 'Trang chủ & khám phá',
                    'Duyệt danh sách quán, lọc theo danh mục, tìm kiếm',
                    'Khách hàng'),
                   (('4', WD_ALIGN_PARAGRAPH.CENTER), 'Chi tiết quán & thực đơn',
                    'Xem thông tin quán, danh sách món ăn, chi tiết món',
                    'Khách hàng'),
                   (('5', WD_ALIGN_PARAGRAPH.CENTER), 'Giỏ hàng',
                    'Thêm/xóa/điều chỉnh số lượng món, xem tổng tiền',
                    'Khách hàng'),
                   (('6', WD_ALIGN_PARAGRAPH.CENTER), 'Đặt hàng',
                    'Nhập địa chỉ giao, xác nhận đơn, theo dõi trạng thái',
                    'Khách hàng'),
                   (('7', WD_ALIGN_PARAGRAPH.CENTER), 'Lịch sử đơn hàng',
                    'Xem các đơn đã đặt và trạng thái xử lý real-time',
                    'Khách hàng'),
                   (('8', WD_ALIGN_PARAGRAPH.CENTER), 'Yêu thích',
                    'Thêm/xóa quán ăn yêu thích, xem danh sách',
                    'Khách hàng'),
                   (('9', WD_ALIGN_PARAGRAPH.CENTER), 'Dashboard chủ quán',
                    'Tổng quan: đơn chờ, doanh thu ngày, số món',
                    'Chủ quán'),
                   (('10', WD_ALIGN_PARAGRAPH.CENTER), 'Quản lý đơn hàng',
                    'Xem, xác nhận, từ chối, hoàn thành đơn real-time',
                    'Chủ quán'),
                   (('11', WD_ALIGN_PARAGRAPH.CENTER), 'Quản lý thực đơn',
                    'Thêm, sửa, xóa món ăn; upload ảnh',
                    'Chủ quán'),
                   (('12', WD_ALIGN_PARAGRAPH.CENTER), 'Thống kê doanh thu',
                    'Biểu đồ doanh thu theo tuần/tháng, số đơn theo ngày',
                    'Chủ quán'),
                   (('13', WD_ALIGN_PARAGRAPH.CENTER), 'Cài đặt quán',
                    'Cập nhật thông tin, ảnh, địa chỉ quán ăn',
                    'Chủ quán'),
               ],
               col_widths=[0.8, 4.0, 6.0, 3.7],
               header_bg='365F91')

    # 3.2 Use case
    h2(doc, '3.2 Sơ đồ use-case')
    body(doc,
         'Phần này mô tả chi tiết các use case chính của hệ thống. Mỗi bảng ứng với '
         'một nhóm chức năng, trong đó cột "Yêu cầu chi tiết" liệt kê các bước xử lý '
         'và điều kiện nghiệp vụ cụ thể.')
    spacer(doc, 4)

    use_case_groups = [
        ('1', 'Xác thực và quản lý người dùng', [
            ('1.1', 'Đăng ký tài khoản',
             'Cho phép người dùng tạo tài khoản mới',
             '– Nhập email hợp lệ, mật khẩu (≥ 6 ký tự)\n– Firebase tạo UID và lưu profile vào Firestore\n– Chuyển về LoginActivity sau khi thành công'),
            ('1.2', 'Đăng nhập',
             'Xác thực người dùng để truy cập hệ thống',
             '– Nhập email và mật khẩu\n– Firebase Auth kiểm tra và trả về token\n– Đọc role từ Firestore → điều hướng phù hợp'),
            ('1.3', 'Đăng xuất',
             'Kết thúc phiên làm việc hiện tại',
             '– Gọi Firebase signOut()\n– Xóa phiên và chuyển về LoginActivity'),
            ('1.4', 'Đặt lại mật khẩu',
             'Gửi email đặt lại mật khẩu qua Firebase',
             '– Nhập email đã đăng ký\n– Firebase gửi link reset đến email người dùng'),
        ]),
        ('2', 'Quản lý hồ sơ cá nhân', [
            ('2.1', 'Xem hồ sơ',
             'Hiển thị thông tin cá nhân',
             '– Tên, email, số điện thoại, ảnh đại diện\n– Đọc document Users từ Firestore theo UID'),
            ('2.2', 'Cập nhật thông tin',
             'Chỉnh sửa tên hiển thị, số điện thoại',
             '– Nhập thông tin mới và lưu lên Firestore\n– Cập nhật cả Firebase Auth displayName'),
            ('2.3', 'Đổi ảnh đại diện',
             'Upload ảnh mới từ thiết bị lên Cloudinary',
             '– Chọn ảnh từ thư viện thiết bị\n– Upload qua CloudinaryHelper, nhận URL\n– Lưu URL vào trường photoURL trong Firestore'),
        ]),
        ('3', 'Trang chủ và khám phá', [
            ('3.1', 'Xem danh sách quán',
             'Hiển thị tất cả quán ăn đang hoạt động',
             '– Query collection Stores (isActive = true)\n– Hiển thị qua RecyclerView với StoreAdapter\n– Lắng nghe real-time qua addSnapshotListener'),
            ('3.2', 'Lọc theo danh mục',
             'Lọc quán ăn theo danh mục món ăn',
             '– Hiển thị danh sách danh mục ở thanh ngang\n– Người dùng chọn danh mục → query Stores theo categoryIds'),
            ('3.3', 'Tìm kiếm quán',
             'Tìm quán ăn theo tên',
             '– Nhập từ khoá vào SearchView\n– Lọc danh sách quán theo tên real-time (client-side)'),
            ('3.4', 'Xem chi tiết quán',
             'Hiển thị thông tin quán và danh sách thực đơn',
             '– Nhấn vào quán → mở StoreDetailActivity\n– Load thông tin quán và danh sách Foods từ Firestore'),
            ('3.5', 'Xem chi tiết món ăn',
             'Hiển thị mô tả đầy đủ, ảnh và giá món',
             '– Nhấn vào món → mở FoodDetailActivity\n– Nút "Thêm vào giỏ" gọi CartManager.addItem()'),
        ]),
        ('4', 'Giỏ hàng và đặt hàng', [
            ('4.1', 'Thêm món vào giỏ',
             'Thêm món ăn vào CartManager (in-memory)',
             '– Kiểm tra quán hiện tại với quán trong giỏ\n– Nếu khác quán: hỏi người dùng có xóa giỏ cũ không\n– Thêm CartItem hoặc tăng số lượng nếu đã có'),
            ('4.2', 'Quản lý số lượng',
             'Tăng/giảm số lượng từng món trong giỏ',
             '– Nhấn nút +/– trên CartFragment\n– Cập nhật CartManager và tính lại tổng tiền'),
            ('4.3', 'Xóa món khỏi giỏ',
             'Xóa một hoặc toàn bộ món trong giỏ',
             '– Nhấn nút xóa trên CartFragment\n– Gọi CartManager.removeItem() hoặc clearCart()'),
            ('4.4', 'Thanh toán đơn hàng',
             'Xác nhận đơn và tạo document trong Firestore',
             '– Nhập địa chỉ giao hàng trong CheckoutActivity\n– Tạo Order document với status = "pending"\n– Xóa giỏ hàng sau khi đặt thành công'),
            ('4.5', 'Xem lịch sử đơn',
             'Hiển thị danh sách đơn hàng đã đặt',
             '– Query Orders theo customerId, sắp xếp theo createdAt\n– Lắng nghe real-time để cập nhật trạng thái'),
        ]),
        ('5', 'Yêu thích', [
            ('5.1', 'Thêm quán yêu thích',
             'Đánh dấu quán ăn là yêu thích',
             '– Nhấn nút tim trên StoreDetailActivity\n– Tạo document trong collection Favorites'),
            ('5.2', 'Xóa quán yêu thích',
             'Bỏ đánh dấu yêu thích',
             '– Nhấn lại nút tim (toggle off)\n– Xóa document tương ứng trong Favorites'),
            ('5.3', 'Xem danh sách yêu thích',
             'Hiển thị tất cả quán đã được yêu thích',
             '– Query Favorites theo userId\n– Load thông tin quán và hiển thị trong FavoritesFragment'),
        ]),
        ('6', 'Chức năng chủ quán', [
            ('6.1', 'Xem dashboard',
             'Hiển thị tổng quan tình hình kinh doanh',
             '– Số đơn đang chờ xử lý\n– Doanh thu hôm nay và tuần này\n– Tổng số món ăn đang bán'),
            ('6.2', 'Quản lý đơn hàng',
             'Xem, xác nhận, hoàn thành hoặc từ chối đơn',
             '– Lắng nghe real-time qua addSnapshotListener\n– Cập nhật trạng thái Order trong Firestore\n– Khách hàng thấy cập nhật ngay lập tức'),
            ('6.3', 'Quản lý thực đơn',
             'Thêm, sửa, xóa món ăn trong thực đơn',
             '– CRUD trên collection Foods\n– Upload ảnh món qua Cloudinary\n– Bật/tắt trạng thái isAvailable'),
            ('6.4', 'Thống kê doanh thu',
             'Xem thống kê doanh thu và số đơn theo thời gian',
             '– Hiển thị doanh thu ngày/tuần/tháng\n– Lọc theo khoảng thời gian\n– Hiển thị tổng doanh thu, số đơn hoàn thành'),
            ('6.5', 'Cài đặt quán',
             'Cập nhật thông tin, ảnh bìa và địa chỉ quán',
             '– Sửa tên quán, địa chỉ, mô tả, giờ mở cửa\n– Upload ảnh bìa qua Cloudinary\n– Lưu vào collection Stores'),
        ]),
    ]

    for grp_num, grp_name, ucs in use_case_groups:
        body(doc, f'Nhóm {grp_num}: {grp_name}', bold=True, indent=False, space_after=4)
        n_ucs = len(ucs)
        t = doc.add_table(rows=1 + n_ucs * 2, cols=4)
        t.style = 'Table Grid'
        for i, h in enumerate(['STT', 'Chức năng', 'Mô tả', 'Yêu cầu chi tiết']):
            set_cell_bg(t.rows[0].cells[i], '365F91')
            fill_cell(t.rows[0].cells[i], h, bold=True,
                      align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
        for ui, (uc_num, uc_name, uc_desc, uc_detail) in enumerate(ucs):
            row_idx = 1 + ui * 2
            for ci, (val, al) in enumerate([
                (uc_num, WD_ALIGN_PARAGRAPH.CENTER),
                (uc_name, WD_ALIGN_PARAGRAPH.LEFT),
                (uc_desc, WD_ALIGN_PARAGRAPH.LEFT),
                (uc_detail, WD_ALIGN_PARAGRAPH.LEFT),
            ]):
                fill_cell(t.rows[row_idx].cells[ci], val, align=al, size=12)
            for ci in range(4):
                fill_cell(t.rows[row_idx + 1].cells[ci], '', size=11)
        for row in t.rows:
            row.cells[0].width = Cm(1.0)
            row.cells[1].width = Cm(3.3)
            row.cells[2].width = Cm(3.7)
            row.cells[3].width = Cm(6.5)
        spacer(doc, 10)

    # 3.3 CSDL
    h2(doc, '3.3 Thiết kế cơ sở dữ liệu')
    body(doc,
         'FoodNow sử dụng Cloud Firestore với 6 collection chính. Mỗi document trong '
         'Firestore là một bản ghi dạng JSON với cấu trúc linh hoạt. Dưới đây là '
         'thiết kế chi tiết từng collection:')
    spacer(doc, 4)

    collections = [
        ('3.3.1 Collection Users',
         'Lưu thông tin người dùng (cả khách hàng và chủ quán).', [
             ('uid', 'string', 'ID người dùng (Document ID)', 'Khóa chính, lấy từ Firebase Auth UID'),
             ('email', 'string', 'Địa chỉ email đăng nhập', 'Bắt buộc, duy nhất'),
             ('displayName', 'string', 'Tên hiển thị của người dùng', 'Bắt buộc'),
             ('photoURL', 'string', 'URL ảnh đại diện trên Cloudinary', 'Tùy chọn'),
             ('phoneNumber', 'string', 'Số điện thoại liên hệ', 'Tùy chọn'),
             ('role', 'string', 'Vai trò: "customer" hoặc "owner"', 'Bắt buộc, mặc định: "customer"'),
             ('storeId', 'string', 'ID quán (chỉ có nếu role = "owner")', 'Tùy chọn, khóa ngoại → Stores'),
             ('createdAt', 'timestamp', 'Thời điểm tạo tài khoản', 'Tự động tạo'),
             ('updatedAt', 'timestamp', 'Thời điểm cập nhật cuối', 'Tự động cập nhật'),
         ]),
        ('3.3.2 Collection Stores',
         'Lưu thông tin các quán ăn.', [
             ('storeId', 'string', 'ID quán (Document ID)', 'Khóa chính, tự động tạo'),
             ('ownerId', 'string', 'UID chủ quán', 'Bắt buộc, khóa ngoại → Users'),
             ('name', 'string', 'Tên quán ăn', 'Bắt buộc'),
             ('description', 'string', 'Mô tả quán', 'Tùy chọn'),
             ('address', 'string', 'Địa chỉ quán ăn', 'Bắt buộc'),
             ('imageURL', 'string', 'URL ảnh bìa quán trên Cloudinary', 'Tùy chọn'),
             ('categoryIds', 'array<string>', 'Mảng ID danh mục phục vụ', 'Tùy chọn'),
             ('openingHours', 'string', 'Giờ mở cửa (VD: "07:00 – 22:00")', 'Tùy chọn'),
             ('rating', 'number', 'Điểm đánh giá trung bình', 'Tùy chọn, mặc định: 0.0'),
             ('isActive', 'boolean', 'Quán có đang mở không', 'Bắt buộc, mặc định: true'),
             ('createdAt', 'timestamp', 'Thời điểm tạo', 'Tự động tạo'),
         ]),
        ('3.3.3 Collection Foods',
         'Lưu thông tin các món ăn trong thực đơn.', [
             ('foodId', 'string', 'ID món ăn (Document ID)', 'Khóa chính, tự động tạo'),
             ('storeId', 'string', 'ID quán sở hữu', 'Bắt buộc, khóa ngoại → Stores'),
             ('name', 'string', 'Tên món ăn', 'Bắt buộc'),
             ('description', 'string', 'Mô tả món ăn', 'Tùy chọn'),
             ('price', 'number', 'Giá món (VNĐ)', 'Bắt buộc, ≥ 0'),
             ('imageURL', 'string', 'URL ảnh món trên Cloudinary', 'Tùy chọn'),
             ('categoryId', 'string', 'ID danh mục', 'Tùy chọn, khóa ngoại → Categories'),
             ('isAvailable', 'boolean', 'Còn hàng hay không', 'Bắt buộc, mặc định: true'),
             ('createdAt', 'timestamp', 'Thời điểm tạo', 'Tự động tạo'),
             ('updatedAt', 'timestamp', 'Thời điểm cập nhật', 'Tự động cập nhật'),
         ]),
        ('3.3.4 Collection Orders',
         'Lưu thông tin đơn hàng.', [
             ('orderId', 'string', 'ID đơn hàng (Document ID)', 'Khóa chính, tự động tạo'),
             ('customerId', 'string', 'UID khách hàng', 'Bắt buộc, khóa ngoại → Users'),
             ('storeId', 'string', 'ID quán nhận đơn', 'Bắt buộc, khóa ngoại → Stores'),
             ('items', 'array<OrderItem>', 'Danh sách món (tên, giá, số lượng)', 'Bắt buộc, ≥ 1 phần tử'),
             ('totalAmount', 'number', 'Tổng tiền đơn hàng (VNĐ)', 'Bắt buộc, ≥ 0'),
             ('deliveryAddress', 'string', 'Địa chỉ giao hàng', 'Bắt buộc'),
             ('status', 'string', 'Trạng thái đơn hàng',
              'Enum: pending / confirmed / completed / cancelled'),
             ('note', 'string', 'Ghi chú của khách hàng', 'Tùy chọn'),
             ('createdAt', 'timestamp', 'Thời điểm đặt hàng', 'Tự động tạo'),
             ('updatedAt', 'timestamp', 'Thời điểm cập nhật trạng thái', 'Tự động cập nhật'),
         ]),
        ('3.3.5 Collection Favorites',
         'Lưu danh sách quán yêu thích của khách hàng.', [
             ('favoriteId', 'string', 'ID bản ghi (Document ID)', 'Khóa chính, tự động tạo'),
             ('userId', 'string', 'UID khách hàng', 'Bắt buộc, khóa ngoại → Users'),
             ('storeId', 'string', 'ID quán yêu thích', 'Bắt buộc, khóa ngoại → Stores'),
             ('createdAt', 'timestamp', 'Thời điểm thêm vào yêu thích', 'Tự động tạo'),
         ]),
        ('3.3.6 Collection Categories',
         'Lưu danh mục món ăn.', [
             ('categoryId', 'string', 'ID danh mục (Document ID)', 'Khóa chính'),
             ('name', 'string', 'Tên danh mục (VD: Lẩu, Cơm, Bún)', 'Bắt buộc, duy nhất'),
             ('imageURL', 'string', 'URL icon danh mục trên Cloudinary', 'Tùy chọn'),
             ('createdAt', 'timestamp', 'Thời điểm tạo', 'Tự động tạo'),
         ]),
    ]

    for col_title, col_desc, fields in collections:
        h3(doc, col_title)
        body(doc, col_desc, space_after=4)
        make_table(doc,
                   headers=['Tên thuộc tính', 'Kiểu dữ liệu', 'Mô tả', 'Ràng buộc'],
                   data_rows=[(f, t, d, r) for f, t, d, r in fields],
                   col_widths=[3.5, 3.0, 4.5, 3.5],
                   header_bg='365F91')

    # 3.4 Luồng điều hướng
    h2(doc, '3.4 Thiết kế luồng điều hướng')
    body(doc,
         'Hệ thống FoodNow có hai luồng điều hướng chính, được phân chia dựa trên '
         'vai trò của người dùng sau khi đăng nhập thành công:')
    spacer(doc, 4)

    body(doc, 'Luồng khách hàng (Customer flow):', bold=True, indent=False, space_after=2)
    for step in [
        'LoginActivity / RegisterActivity → (xác thực thành công, role = "customer") → MainActivity',
        'MainActivity (BottomNavigationView 4 tab): HomeFragment | OrdersFragment | FavoritesFragment | ProfileFragment',
        'HomeFragment → (nhấn vào quán) → StoreDetailActivity → (nhấn vào món) → FoodDetailActivity',
        'FoodDetailActivity / StoreDetailActivity → (nhấn "Thêm vào giỏ") → CartFragment',
        'CartFragment → (nhấn "Đặt hàng") → CheckoutActivity → (xác nhận) → cập nhật OrdersFragment',
    ]:
        bullet(doc, step)
    spacer(doc, 4)

    body(doc, 'Luồng chủ quán (Store Owner flow):', bold=True, indent=False, space_after=2)
    for step in [
        'LoginActivity → (đăng nhập thành công, role = "owner") → StoreOwnerActivity',
        'StoreOwnerActivity (BottomNavigationView 4 tab): DashboardFragment | StoreOrdersFragment | ManageFoodsFragment | OwnerSettingsFragment',
        'ManageFoodsFragment → (thêm/sửa món) → AddEditFoodActivity',
        'OwnerSettingsFragment → EditStoreActivity (sửa thông tin quán) / EditProfileActivity (sửa hồ sơ)',
        'DashboardFragment → StatsFragment (xem biểu đồ doanh thu chi tiết)',
    ]:
        bullet(doc, step)

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# CHƯƠNG 4: XÂY DỰNG ỨNG DỤNG
# ═══════════════════════════════════════════════════════════════════

def build_chapter4(doc):
    h1(doc, 'CHƯƠNG 4: XÂY DỰNG ỨNG DỤNG')

    # 4.1 Môi trường
    h2(doc, '4.1 Môi trường phát triển')

    h3(doc, '4.1.1 Android Studio')
    body(doc,
         'Android Studio là IDE chính thức cho phát triển ứng dụng Android, được '
         'Google phát triển dựa trên IntelliJ IDEA. Android Studio cung cấp tích hợp '
         'đầy đủ với Android SDK, Gradle build system, AVD Manager (máy ảo Android), '
         'Logcat, Layout Inspector và các công cụ debug/profile.')
    body(doc,
         'FoodNow được phát triển trên Android Studio Electric Eel trở lên, với '
         'compileSdk 35, minSdk 24 và targetSdk 35. Build system sử dụng Gradle '
         'với Kotlin DSL (build.gradle.kts) thay vì Groovy DSL truyền thống.')

    h3(doc, '4.1.2 Git và GitHub')
    body(doc,
         'Git là hệ thống quản lý phiên bản phân tán, được nhóm sử dụng để quản lý '
         'mã nguồn trong suốt quá trình phát triển. Mỗi thành viên làm việc trên '
         'nhánh riêng và merge vào nhánh main sau khi hoàn thiện module.')
    body(doc,
         'Repository được lưu trữ trên GitHub tại: github.com/minnhi09/FoodNow-k47pm. '
         'GitHub cũng được dùng để theo dõi phân công công việc và review code giữa '
         'các thành viên trong nhóm.')

    h3(doc, '4.1.3 Figma')
    body(doc,
         'Figma là công cụ thiết kế giao diện UI/UX trực tuyến được nhóm sử dụng '
         'để thiết kế wireframe và prototype giao diện trước khi triển khai code. '
         'Thiết kế tuân theo Material Design 3 với bảng màu chủ đạo cam (#FF6B35) '
         'phù hợp với nhận diện thương hiệu ứng dụng đặt đồ ăn.')

    # 4.2 Cấu trúc dự án
    h2(doc, '4.2 Cấu trúc dự án')
    body(doc,
         'Dự án được tổ chức theo kiến trúc MVVM với cấu trúc package rõ ràng. '
         'Bảng dưới đây mô tả vai trò của từng package:')
    spacer(doc, 4)

    make_table(doc,
               headers=['Package', 'Mô tả'],
               data_rows=[
                   ('models',
                    'Data classes tương ứng với Firestore: User, Store, Food, Order, Favorite, Category, CartItem. '
                    'Mỗi class cần empty constructor và getters/setters để Firestore deserialize.'),
                   ('repositories',
                    'Tầng truy cập dữ liệu Firebase: AuthRepository, UserRepository, StoreRepository, '
                    'FoodRepository, OrderRepository, FavoriteRepository, CategoryRepository.'),
                   ('viewmodels',
                    'Logic nghiệp vụ và LiveData: AuthViewModel, ProfileViewModel, HomeViewModel, '
                    'StoreDetailViewModel, CheckoutViewModel, OrdersViewModel, FavoritesViewModel.'),
                   ('fragments',
                    'Màn hình tab: HomeFragment, CartFragment, OrdersFragment, FavoritesFragment, '
                    'ProfileFragment, DashboardFragment, ManageFoodsFragment, StoreOrdersFragment, '
                    'OwnerSettingsFragment, StatsFragment.'),
                   ('activities',
                    'Màn hình độc lập: LoginActivity, RegisterActivity, MainActivity, '
                    'StoreDetailActivity, FoodDetailActivity, CheckoutActivity, StoreOwnerActivity, '
                    'AddEditFoodActivity, EditStoreActivity, EditProfileActivity, ImageAdminActivity.'),
                   ('adapters',
                    'RecyclerView adapters: StoreAdapter, FoodAdapter, CartAdapter, OrderAdapter, '
                    'FavoriteAdapter, ManageFoodAdapter, StoreOrderAdapter, CategoryAdapter.'),
                   ('utils',
                    'Lớp tiện ích: CloudinaryHelper (xử lý upload ảnh), CartManager (Singleton quản lý giỏ).'),
               ],
               col_widths=[3.5, 11.0],
               header_bg='365F91')

    # 4.3 Modules
    h2(doc, '4.3 Xây dựng các module chính')

    h3(doc, '4.3.1 Module xác thực')
    body(doc,
         'Module xác thực bao gồm LoginActivity và RegisterActivity, sử dụng '
         'AuthViewModel và AuthRepository để giao tiếp với Firebase Authentication.')
    body(doc, 'Luồng đăng ký:', space_after=2)
    for s in [
        'Người dùng nhập email và mật khẩu trong RegisterActivity.',
        'AuthRepository.register() gọi createUserWithEmailAndPassword() của Firebase Auth.',
        'Sau khi tạo tài khoản thành công, tạo document trong collection Users với role = "customer".',
        'Hiển thị thông báo thành công và chuyển về LoginActivity.',
    ]:
        p = doc.add_paragraph(style='List Paragraph')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.add_run('– ' + s)
    spacer(doc, 4)
    body(doc, 'Luồng đăng nhập:', space_after=2)
    for s in [
        'Người dùng nhập email và mật khẩu trong LoginActivity.',
        'AuthRepository.login() gọi signInWithEmailAndPassword().',
        'Đọc document Users để xác định role của người dùng.',
        'Nếu role = "customer": chuyển đến MainActivity. Nếu role = "owner": chuyển đến StoreOwnerActivity.',
    ]:
        p = doc.add_paragraph(style='List Paragraph')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.add_run('– ' + s)

    h3(doc, '4.3.2 Module trang chủ và quán ăn')
    body(doc,
         'Module này bao gồm HomeFragment và StoreDetailActivity, sử dụng '
         'HomeViewModel, StoreDetailViewModel kết hợp với StoreRepository và '
         'FoodRepository.')
    for item in [
        'HomeFragment: hiển thị banner danh mục (RecyclerView ngang) và danh sách quán ăn (RecyclerView dọc). '
        'Sử dụng addSnapshotListener() để lắng nghe cập nhật real-time từ Firestore.',
        'StoreDetailActivity: nhận storeId qua Intent, load thông tin quán và danh sách thực đơn. '
        'Người dùng có thể nhấn vào món ăn để xem FoodDetailActivity hoặc thêm trực tiếp vào giỏ.',
        'FoodDetailActivity: hiển thị ảnh, mô tả đầy đủ và giá món ăn, cùng với nút "Thêm vào giỏ" '
        'gọi CartManager.addItem().',
    ]:
        bullet(doc, item)

    h3(doc, '4.3.3 Module giỏ hàng và đặt hàng')
    body(doc,
         'CartManager là Singleton quản lý giỏ hàng trong bộ nhớ (in-memory), '
         'không lưu vào Firestore. Giỏ hàng chỉ chứa món từ một quán tại một thời '
         'điểm — nếu thêm món từ quán khác, người dùng được hỏi có muốn xóa giỏ cũ không.')
    for item in [
        'CartFragment: hiển thị danh sách CartItem với nút +/– điều chỉnh số lượng và nút "Đặt hàng" '
        'mở CheckoutActivity.',
        'CheckoutActivity: form nhập địa chỉ giao hàng và ghi chú. Khi xác nhận, tạo Order document '
        'trong Firestore với status = "pending".',
        'OrdersFragment: query Orders theo customerId, lắng nghe real-time để hiển thị trạng thái '
        'đơn hàng được cập nhật tức thì khi chủ quán xử lý.',
    ]:
        bullet(doc, item)

    h3(doc, '4.3.4 Module yêu thích')
    body(doc,
         'FavoritesFragment hiển thị danh sách quán mà người dùng đã đánh dấu yêu thích. '
         'FavoriteRepository thực hiện thêm/xóa và query collection Favorites theo userId. '
         'Khi nhấn vào quán yêu thích, ứng dụng điều hướng đến StoreDetailActivity để '
         'xem thực đơn và đặt hàng.')

    h3(doc, '4.3.5 Module chủ quán')
    body(doc,
         'StoreOwnerActivity là Activity dành riêng cho chủ quán, có cấu trúc tương tự '
         'MainActivity với BottomNavigationView 4 tab:')
    for tab, desc in [
        ('DashboardFragment', 'hiển thị tổng quan số đơn chờ, doanh thu ngày và tổng số món ăn.'),
        ('StoreOrdersFragment', 'danh sách đơn hàng real-time với nút xác nhận/từ chối/hoàn thành.'),
        ('ManageFoodsFragment', 'danh sách thực đơn với nút thêm/sửa/xóa, liên kết AddEditFoodActivity.'),
        ('OwnerSettingsFragment', 'liên kết đến StatsFragment (biểu đồ), EditStoreActivity và EditProfileActivity.'),
    ]:
        p = doc.add_paragraph(style='List Paragraph')
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after = Pt(2)
        p.add_run(tab + ': ').bold = True
        p.add_run(desc)
    spacer(doc, 4)
    body(doc,
         'AddEditFoodActivity hỗ trợ cả chế độ thêm mới và chỉnh sửa món ăn. '
         'Ảnh món được upload lên Cloudinary qua CloudinaryHelper, URL nhận về '
         'được lưu vào trường imageURL trong Firestore.')

    # 4.4 Phân công
    h2(doc, '4.4 Phân công công việc')
    body(doc,
         'Dự án được phân chia thành 3 phần tương ứng với 3 thành viên, dựa trên '
         'tính độc lập và logic của từng module. Thứ tự thực hiện: TV1 hoàn thành '
         'phần Auth và setup trước, sau đó TV2 và TV3 làm song song.')
    spacer(doc, 6)

    work_data = [
        # TV1 — Đinh Thị Mai Lành
        ('TV1', 'Đinh Thị Mai Lành\n(2312660)',
         'Setup & Config', 'AndroidManifest.xml, build.gradle.kts, Firebase config, Cloudinary config, activity_main.xml, bottom_nav_menu.xml'),
        ('', '', 'Models', 'User.java, Category.java'),
        ('', '', 'Repositories', 'AuthRepository.java, UserRepository.java, CategoryRepository.java'),
        ('', '', 'ViewModels', 'AuthViewModel.java, ProfileViewModel.java'),
        ('', '', 'Activities (Auth)', 'LoginActivity.java, RegisterActivity.java'),
        ('', '', 'Activities (Owner)', 'StoreOwnerActivity.java, AddEditFoodActivity.java, EditStoreActivity.java, EditProfileActivity.java, ImageAdminActivity.java, FoodDetailActivity.java'),
        ('', '', 'Fragments', 'ProfileFragment.java, DashboardFragment.java, ManageFoodsFragment.java, StoreOrdersFragment.java, OwnerSettingsFragment.java, StatsFragment.java'),
        ('', '', 'Adapters', 'CategoryAdapter.java, ManageFoodAdapter.java, StoreOrderAdapter.java'),
        ('', '', 'Utils', 'CloudinaryHelper.java'),
        ('', '', 'MainActivity', 'MainActivity.java (setup BottomNavigationView + 4 tab điều hướng)'),
        # TV2 — Võ Thị Minh Ân
        ('TV2', 'Võ Thị Minh Ân\n(2312567)',
         'Models', 'Store.java, Food.java'),
        ('', '', 'Repositories', 'StoreRepository.java, FoodRepository.java'),
        ('', '', 'ViewModels', 'HomeViewModel.java, StoreDetailViewModel.java'),
        ('', '', 'Activities', 'StoreDetailActivity.java'),
        ('', '', 'Fragments', 'HomeFragment.java'),
        ('', '', 'Adapters', 'StoreAdapter.java, FoodAdapter.java'),
        ('', '', 'Layouts', 'fragment_home.xml, activity_store_detail.xml, item_store.xml, item_food.xml'),
        # TV3 — Trương Võ Trọng Nhân
        ('TV3', 'Trương Võ Trọng Nhân\n(2312708)',
         'Models', 'CartItem.java, Order.java (+ OrderItem inner class), Favorite.java'),
        ('', '', 'Repositories', 'OrderRepository.java, FavoriteRepository.java'),
        ('', '', 'ViewModels', 'CheckoutViewModel.java, OrdersViewModel.java, FavoritesViewModel.java'),
        ('', '', 'Activities', 'CheckoutActivity.java'),
        ('', '', 'Fragments', 'CartFragment.java, OrdersFragment.java, FavoritesFragment.java'),
        ('', '', 'Adapters', 'CartAdapter.java, OrderAdapter.java, FavoriteAdapter.java'),
        ('', '', 'Utils', 'CartManager.java (Singleton quản lý giỏ hàng in-memory)'),
        ('', '', 'Layouts', 'fragment_cart.xml, activity_checkout.xml, fragment_orders.xml, fragment_favorites.xml, item_cart.xml, item_order.xml, item_favorite.xml'),
    ]

    n_rows = 1 + len(work_data)
    t = doc.add_table(rows=n_rows, cols=4)
    t.style = 'Table Grid'
    for i, h in enumerate(['Thành viên', 'Họ tên / MSSV', 'Layer', 'File thực hiện']):
        set_cell_bg(t.rows[0].cells[i], '365F91')
        fill_cell(t.rows[0].cells[i], h, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    for ri, (tv, name, layer, files) in enumerate(work_data):
        row = t.rows[ri + 1]
        fill_cell(row.cells[0], tv, align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
        fill_cell(row.cells[1], name, size=11)
        fill_cell(row.cells[2], layer, bold=(layer != ''), size=11)
        fill_cell(row.cells[3], files, size=11)
    for row in t.rows:
        row.cells[0].width = Cm(1.5)
        row.cells[1].width = Cm(3.5)
        row.cells[2].width = Cm(2.5)
        row.cells[3].width = Cm(7.0)
    spacer(doc, 8)

    # Summary table
    body(doc, 'Bảng tổng hợp tỷ lệ đóng góp:', bold=True, indent=False, space_after=4)
    make_table(doc,
               headers=['Thành viên', 'Họ tên', 'MSSV', 'Phần thực hiện', 'Tỷ lệ'],
               data_rows=[
                   ('TV1', 'Đinh Thị Mai Lành', '2312660',
                    'Setup, Auth, Profile, toàn bộ hệ thống chủ quán, FoodDetail', '~55%'),
                   ('TV2', 'Võ Thị Minh Ân', '2312567',
                    'Trang chủ, Quán ăn, Thực đơn', '~25%'),
                   ('TV3', 'Trương Võ Trọng Nhân', '2312708',
                    'Giỏ hàng, Đặt hàng, Yêu thích, CartManager', '~20%'),
               ],
               col_widths=[1.5, 3.5, 2.0, 6.0, 1.5],
               header_bg='365F91')

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# CHƯƠNG 5: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN
# ═══════════════════════════════════════════════════════════════════

def build_chapter5(doc):
    h1(doc, 'CHƯƠNG 5: KẾT LUẬN VÀ HƯỚNG PHÁT TRIỂN')

    h2(doc, '5.1 Kết quả đạt được')
    body(doc,
         'Sau quá trình thực hiện đề tài, nhóm đã xây dựng thành công ứng dụng '
         'Android FoodNow với đầy đủ các chức năng đề ra. Cụ thể, hệ thống đã '
         'đạt được các kết quả sau:')
    for item in [
        'Xây dựng hoàn chỉnh ứng dụng Android FoodNow theo kiến trúc MVVM, sử dụng Java và XML Views.',
        'Tích hợp thành công Firebase Authentication với đầy đủ chức năng: đăng ký, đăng nhập, '
        'đăng xuất và đặt lại mật khẩu.',
        'Xây dựng đầy đủ luồng đặt hàng: duyệt quán → xem thực đơn → thêm vào giỏ → '
        'thanh toán → theo dõi đơn hàng theo thời gian thực.',
        'Xây dựng hệ thống chủ quán hoàn chỉnh: dashboard thống kê, xử lý đơn hàng '
        'real-time, quản lý thực đơn CRUD và thống kê doanh thu.',
        'Tích hợp Cloudinary để upload và hiển thị ảnh món ăn, ảnh quán và ảnh đại diện '
        'người dùng qua Glide.',
        'Thiết kế giao diện theo Material Design 3 với bố cục BottomNavigationView, '
        'hỗ trợ Android 7.0 (API 24) trở lên.',
        'Triển khai Firestore Security Rules để bảo vệ dữ liệu theo vai trò người dùng.',
        'Phân chia công việc hiệu quả, nhóm hoàn thành đúng tiến độ và tích lũy kinh '
        'nghiệm thực tế về phát triển ứng dụng Android.',
    ]:
        bullet(doc, item)
    spacer(doc, 4)

    body(doc,
         'Bên cạnh kết quả đạt được, đề tài vẫn còn một số hạn chế:')
    for item in [
        'Hệ thống thanh toán chỉ ở mức mock (giả lập), chưa tích hợp cổng thanh toán thực.',
        'Chưa có hệ thống thông báo đẩy (push notification) khi trạng thái đơn hàng thay đổi.',
        'Chưa tích hợp bản đồ (Google Maps) để hiển thị vị trí quán và tính khoảng cách.',
        'Chưa có chức năng đánh giá và bình luận quán ăn sau khi đặt hàng.',
        'Chưa có cơ chế xử lý đồng thời nhiều người dùng truy cập cùng lúc (concurrency).',
    ]:
        bullet(doc, item)

    h2(doc, '5.2 Hướng phát triển')
    body(doc,
         'Trong tương lai, nhóm đề xuất các hướng phát triển tiếp theo nhằm '
         'hoàn thiện hơn ứng dụng FoodNow:')
    for item in [
        'Tích hợp cổng thanh toán thực (VNPay, MoMo, ZaloPay) để xử lý giao dịch an toàn.',
        'Triển khai Firebase Cloud Messaging (FCM) để gửi thông báo đẩy đến khách hàng '
        'khi trạng thái đơn hàng thay đổi.',
        'Tích hợp Google Maps API để hiển thị vị trí quán trên bản đồ, tính khoảng cách '
        'và hỗ trợ điều hướng.',
        'Xây dựng hệ thống đánh giá và bình luận quán ăn, tính điểm rating tự động.',
        'Thêm chức năng khuyến mãi và mã giảm giá cho từng quán ăn.',
        'Chuyển sang Jetpack Compose để xây dựng giao diện hiện đại và dễ bảo trì hơn.',
        'Tối ưu hoá hiệu năng: lazy loading, phân trang (pagination) cho danh sách lớn.',
    ]:
        bullet(doc, item)

    page_break(doc)


# ═══════════════════════════════════════════════════════════════════
# TÀI LIỆU THAM KHẢO
# ═══════════════════════════════════════════════════════════════════

def build_tai_lieu_tham_khao(doc):
    h1(doc, 'TÀI LIỆU THAM KHẢO')
    refs = [
        'Android Developer Documentation. (2024). Android Jetpack — ViewModel & LiveData. '
        'https://developer.android.com/topic/libraries/architecture/viewmodel',
        'Firebase. (2024). Firebase Authentication — Android. '
        'https://firebase.google.com/docs/auth/android/start',
        'Firebase. (2024). Cloud Firestore — Android. '
        'https://firebase.google.com/docs/firestore',
        'Cloudinary. (2024). Android SDK Documentation. '
        'https://cloudinary.com/documentation/android_integration',
        'Bumptech. (2024). Glide — Image Loading Framework for Android. '
        'https://github.com/bumptech/glide',
        'Google. (2024). Material Design 3 for Android. '
        'https://m3.material.io/',
        'Google. (2024). Android RecyclerView. '
        'https://developer.android.com/guide/topics/ui/layout/recyclerview',
        'GitHub. (2024). FoodNow Repository — Nhóm 9. '
        'https://github.com/minnhi09/FoodNow-k47pm',
    ]
    for i, ref in enumerate(refs, 1):
        p = doc.add_paragraph(style='Normal')
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.left_indent = Cm(0.5)
        p.add_run(f'[{i}] {ref}')


# ═══════════════════════════════════════════════════════════════════
# MAIN
# ═══════════════════════════════════════════════════════════════════

def build_report():
    doc = Document()

    # Page margins theo mẫu: trái 3cm, phải 2cm, trên 2cm, dưới 3.17cm
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(3.17)
        section.left_margin = Cm(3.0)
        section.right_margin = Cm(2.0)

    setup_styles(doc)

    build_cover(doc)
    build_nhan_xet(doc)
    build_loi_cam_on(doc)
    build_de_cuong(doc)
    build_muc_luc(doc)
    build_chapter1(doc)
    build_chapter2(doc)
    build_chapter3(doc)
    build_chapter4(doc)
    build_chapter5(doc)
    build_tai_lieu_tham_khao(doc)

    out = os.path.join(
        os.path.dirname(os.path.dirname(os.path.abspath(__file__))),
        'BaoCaoDoAn_FoodNow_Nhom9.docx'
    )
    doc.save(out)
    print(f'✅ Đã tạo file báo cáo: {out}')


if __name__ == '__main__':
    build_report()
